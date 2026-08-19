# -*- coding: utf-8 -*-
"""Generate COMP8043041 Assignment II report (DOCX + PDF) from CSVs and figures."""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(r"E:\Work-Project\machine-learning-tr")
OUT_DOCX = ROOT / "COMP8043041_Laporan_Assignment_II.docx"
OUT_PDF = ROOT / "COMP8043041_Laporan_Assignment_II.pdf"

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = "1E3A5F"
HEADER_FILL = "1E3A5F"
ROW_ALT = "F3F6FA"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, size=9, color=None, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def add_toc_field(paragraph, switches=r'TOC \o "1-3" \h \z \u'):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = switches
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Perbarui daftar isi di Word (References → Update Table)."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9, color=RGBColor(255, 255, 255))
        shade_cell(cell, HEADER_FILL)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            align = "left" if c_i == 0 else "center"
            set_cell_text(cell, val, size=9, align=align)
            if r_i % 2 == 1:
                shade_cell(cell, ROW_ALT)
    for row in table.rows:
        prevent_row_split(row)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True, color=NAVY)
    return p


def body(doc, text, *, first_indent=True, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.left_indent = Cm(1.25 + 0.5 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_figure(doc, path, caption_text, width_cm=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    caption(doc, caption_text)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def fmt3(x):
    return f"{float(x):.3f}"


def fmt4(x):
    return f"{float(x):.4f}"


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    for i, size in ((1, 16), (2, 14), (3, 12)):
        st = styles[f"Heading {i}"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.paragraph_format.space_before = Pt(14 if i == 1 else 10)
        st.paragraph_format.space_after = Pt(8)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("COMP8043041 – Machine Learning  |  Assignment II")
    set_run_font(run, size=9, italic=True, color=NAVY)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("BINUS University Graduate Program  ·  Halaman ")
    set_run_font(r1, size=9, color=NAVY)
    add_page_number(fp)


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BINUS UNIVERSITY")
    set_run_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GRADUATE PROGRAM")
    set_run_font(r, size=14, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("People. Innovation. Excellence.")
    set_run_font(r, size=11, italic=True, color=RGBColor(0x5B, 0x6B, 0x7C))

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("—" * 28)
    set_run_font(r, size=12, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMP8043041 — Machine Learning")
    set_run_font(r, size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Assignment II")
    set_run_font(r, size=12)

    for _ in range(1):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(
        "Prediksi Kegagalan Mesin Industri Menggunakan\n"
        "Random Forest pada Dataset Tidak Seimbang AI4I 2020"
    )
    set_run_font(r, size=18, bold=True, color=NAVY)

    for _ in range(2):
        doc.add_paragraph()

    meta = [
        ("Mata Kuliah", "COMP8043041 – Machine Learning"),
        ("Jenis Tugas", "Assignment II — Laporan Eksperimen"),
        ("Dataset", "AI4I 2020 Predictive Maintenance (Matzka, 2020)"),
        ("Model usulan", "Random Forest (regularized / RF-dominated)"),
        ("Penyusun", "Achmad Rasyid"),
        ("Email", "achmad.ramadhan003@binus.ac.id"),
        ("Repositori", "https://github.com/teh05/machine-learning-tr.git"),
        ("Tahun", "2026"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        set_cell_text(table.rows[i].cells[0], k, bold=True, size=11, align="left", color=NAVY)
        set_cell_text(table.rows[i].cells[1], v, size=11, align="left")
        table.rows[i].cells[0].width = Cm(4.2)
        table.rows[i].cells[1].width = Cm(11.5)

    doc.add_page_break()


def add_preface_lists(doc):
    heading(doc, "Daftar Isi", 1)
    p = doc.add_paragraph()
    add_toc_field(p)
    body(
        doc,
        "Daftar isi di atas diperbarui otomatis saat berkas dibuka di Microsoft Word "
        "(klik kanan → Update Field). Struktur laporan mengikuti rubrik Assignment II: "
        "latar belakang, rumusan masalah, deskripsi data, EDA, pra-pemrosesan, pemilihan model, "
        "pelatihan, penyetelan hiperparameter, evaluasi, analisis, serta saran pengembangan.",
        first_indent=False,
    )

    heading(doc, "Daftar Gambar", 1)
    figures = [
        "Gambar 1. Distribusi kelas target Machine failure dan frekuensi per failure mode",
        "Gambar 2. Distribusi fitur sensor per kelas kegagalan (histogram + KDE)",
        "Gambar 3. Boxplot fitur sensor: Normal versus Machine failure",
        "Gambar 4. Heatmap korelasi fitur sensor terhadap Machine failure",
        "Gambar 5. Scatter plot zona kegagalan: Power zone dan Strain zone",
        "Gambar 6. Penanganan outlier: sebelum versus sesudah Winsorization",
        "Gambar 7. Distribusi kelas training sebelum versus sesudah SMOTE",
        "Gambar 8. KDE fitur turunan Power, Strain, dan Temp_diff",
        "Gambar 9. Flowchart pipeline eksperimen predictive maintenance",
        "Gambar 10. Perbandingan F1 sebelum versus sesudah hyperparameter tuning",
        "Gambar 11. Learning curve Train versus CV (XGBoost, Random Forest, LR-Poly)",
        "Gambar 12. Confusion matrix model utama pada test set",
        "Gambar 13. Ranking model utama (F1, Recall, Precision)",
        "Gambar 14. Ranking model extended (konfigurasi bawaan)",
        "Gambar 15. Perbandingan keseluruhan model pada test set",
        "Gambar 16. Sensitivitas rasio SMOTE pada Random Forest",
        "Gambar 17. Kurva ROC model utama pada test set",
        "Gambar 18. SHAP summary Random Forest",
        "Gambar 19. Ranking kepentingan fitur berdasarkan mean |SHAP|",
    ]
    for item in figures:
        bullet(doc, item)

    heading(doc, "Daftar Tabel", 1)
    tables = [
        "Tabel 1. Ringkasan penelitian terdahulu dan research gap",
        "Tabel 2. Informasi dataset AI4I 2020",
        "Tabel 3. Deskripsi fitur dataset",
        "Tabel 4. Ringkasan kualitas data",
        "Tabel 5. Statistik deskriptif fitur sensor numerik",
        "Tabel 6. Distribusi kelas target Machine failure",
        "Tabel 7. Frekuensi failure mode",
        "Tabel 8. Kolom yang dihapus dan alasannya",
        "Tabel 9. Feature engineering domain",
        "Tabel 10. Encoding fitur kategorikal",
        "Tabel 11. Pembagian data train/validasi/test",
        "Tabel 12. Distribusi kelas sebelum dan sesudah SMOTE (training)",
        "Tabel 13. Justifikasi peran model",
        "Tabel 14. Metrik evaluasi dan prioritas konteks manufaktur",
        "Tabel 15. Ringkasan protokol pelatihan",
        "Tabel 16. Hasil hyperparameter tuning (GridSearchCV + PredefinedSplit)",
        "Tabel 17. F1 sebelum versus sesudah tuning",
        "Tabel 18. Diagnosa gap Train–Validation F1",
        "Tabel 19. Stabilitas Repeated Stratified K-Fold 5×3",
        "Tabel 20. Hasil test model utama (F1@0,5)",
        "Tabel 21. Ranking model extended (default)",
        "Tabel 22. Sensitivitas rasio SMOTE pada Random Forest",
        "Tabel 23. Uji McNemar pada prediksi test",
        "Tabel 24. Ranking kepentingan fitur SHAP (Random Forest)",
    ]
    for item in tables:
        bullet(doc, item)
    doc.add_page_break()


def chapter_1(doc):
    heading(doc, "Bab 1. Pendahuluan", 1)
    heading(doc, "1.1 Latar Belakang Masalah", 2)
    body(
        doc,
        "Keberlanjutan operasional mesin produksi merupakan isu strategis dalam konteks Industri 4.0. "
        "Kegagalan mesin yang tidak terjadwal (unplanned downtime) menimbulkan biaya perbaikan darurat, "
        "penghentian lini produksi, dan gangguan rantai pasok. Pendekatan perawatan konvensional—reactive "
        "maintenance dan preventive maintenance—memiliki keterbatasan: pendekatan reaktif menanggung "
        "konsekuensi kerusakan yang telah terjadi, sedangkan pendekatan preventif berisiko melakukan "
        "intervensi pada komponen yang masih layak operasi (Çınar et al., 2020; Abidi et al., 2022).",
    )
    body(
        doc,
        "Predictive maintenance (PdM) berbasis machine learning dikembangkan untuk memperkirakan risiko "
        "kegagalan berdasarkan data sensor operasional sehingga intervensi dapat dijadwalkan lebih tepat "
        "(Taoufyq et al., 2025; Serradilla et al., 2022). Dalam praktik industri, kejadian kegagalan bersifat "
        "jarang sehingga membentuk klasifikasi tidak seimbang. Kondisi tersebut dapat menghasilkan accuracy "
        "paradox: akurasi tinggi tanpa kemampuan deteksi kelas kegagalan yang memadai (Yang & Iqbal, 2025; "
        "Alnahhal et al., 2026).",
    )
    body(
        doc,
        "Penelitian ini menempatkan prediksi kegagalan mesin sebagai klasifikasi biner pada dataset AI4I 2020 "
        "Predictive Maintenance (Matzka, 2020). Dataset tersebut menyediakan label Machine failure beserta "
        "moda kegagalan terkait dan banyak digunakan sebagai benchmark akademik PdM tabular. Fokus diarahkan "
        "pada Random Forest sebagai model usulan untuk dasar sistem peringatan dini, dengan penekanan pada "
        "keseimbangan Recall–Precision serta interpretabilitas fitur (Breiman, 2001; Brito et al., 2024; "
        "Kareem, 2024).",
    )

    heading(doc, "1.2 Rumusan Masalah", 2)
    body(doc, "Berdasarkan latar belakang tersebut, penelitian ini merumuskan masalah sebagai berikut.", first_indent=False)
    bullet(
        doc,
        "Parameter sensor dan fitur turunan manakah yang paling berkontribusi terhadap prediksi Machine failure pada dataset AI4I 2020?",
    )
    bullet(
        doc,
        "Bagaimana membangun model klasifikasi biner yang andal untuk prediksi kegagalan mesin pada data yang sangat tidak seimbang?",
    )
    bullet(
        doc,
        "Bagaimana perbandingan kinerja Random Forest sebagai model usulan terhadap XGBoost, Logistic Regression–Polynomial, Gradient Boosting, serta model extended (termasuk LightGBM) pada protokol evaluasi yang sama?",
    )
    bullet(
        doc,
        "Failure mode manakah yang paling sering terjadi, dan apa implikasinya bagi kebijakan predictive maintenance?",
    )

    heading(doc, "1.3 Tujuan dan Ruang Lingkup", 2)
    heading(doc, "Tujuan penelitian", 3)
    bullet(doc, "Mengidentifikasi parameter sensor mesin (suhu, torsi, kecepatan, keausan alat) dan fitur turunan yang paling berkontribusi terhadap kegagalan.")
    bullet(doc, "Membangun model prediksi klasifikasi biner: apakah mesin akan mengalami kegagalan atau beroperasi normal.")
    bullet(doc, "Membandingkan performa lebih dari satu algoritma machine learning dalam menangani data imbalanced.")
    bullet(doc, "Mengidentifikasi moda kegagalan yang paling sering terjadi sebagai dasar kebijakan maintenance berbasis risiko.")

    heading(doc, "Ruang lingkup", 3)
    bullet(doc, "Dataset: AI4I 2020 Predictive Maintenance Dataset (UCI / Kaggle; Matzka, 2020).")
    bullet(doc, "Target primer: Machine failure (biner; 1 = gagal, 0 = normal).")
    bullet(doc, "Target sekunder (deskriptif): moda kegagalan TWF, HDF, PWF, OSF, dan RNF.")
    bullet(doc, "Model utama: Random Forest (usulan), XGBoost (pembanding), LR–Polynomial (baseline), Gradient Boosting (pembanding tambahan).")
    bullet(doc, "Model extended: LightGBM, CatBoost, ANN/MLP, Decision Tree, AdaBoost, Naive Bayes, SVM, dan KNN pada konfigurasi bawaan.")
    bullet(doc, "Metrik: Accuracy, Precision, Recall, F1-score, dan ROC-AUC, dengan F1@0,5 sebagai kriteria keputusan.")

    heading(doc, "1.4 Tinjauan Penelitian Terdahulu dan Research Gap", 2)
    body(
        doc,
        "Subbab ini meninjau penelitian-penelitian yang relevan untuk menempatkan studi pada konteks predictive "
        "maintenance berbasis machine learning, khususnya pada data manufaktur yang tidak seimbang. Referensi yang "
        "digunakan dibatasi pada artikel jurnal periode 2020–2026 agar tetap mutakhir dan sesuai dengan fokus "
        "perkembangan riset terbaru pada PdM.",
    )
    heading(doc, "Tinjauan pustaka", 3)
    body(
        doc,
        "Pada level umum, literatur menegaskan bahwa predictive maintenance telah berkembang menjadi salah satu "
        "aplikasi utama machine learning dalam manufaktur cerdas. Çınar et al. (2020) memetakan penggunaan berbagai "
        "algoritma seperti Random Forest, XGBoost, SVM, dan neural network untuk mendukung keberlanjutan operasi "
        "pada Industri 4.0. Kajian yang lebih khusus oleh Serradilla et al. (2022) memperlihatkan bahwa meskipun "
        "model deep learning mulai banyak digunakan, model tabular berbasis ensemble masih sangat relevan karena "
        "lebih mudah diterapkan pada data sensor terstruktur dan sering lebih efisien secara komputasi.",
    )
    body(
        doc,
        "Dalam konteks operasional, Abidi et al. (2022) menekankan bahwa keberhasilan PdM tidak cukup diukur dari "
        "akurasi semata, tetapi juga dari kemampuannya mencegah downtime dan mendukung pengambilan keputusan "
        "maintenance. Azari et al. (2023) menambahkan bahwa salah satu keterbatasan utama penelitian PdM adalah "
        "kesenjangan antara data benchmark dan data sensor riil, sehingga isu generalisasi, drift, dan transfer ke "
        "lingkungan industri nyata harus diperhatikan sejak tahap perancangan eksperimen.",
    )
    body(
        doc,
        "Pada sisi pemilihan algoritma, Kareem (2024) menunjukkan bahwa Random Forest dan XGBoost merupakan dua "
        "kandidat kuat untuk kasus PdM tabular karena sama-sama mampu menangkap hubungan nonlinier antar fitur. "
        "Namun, penelitian tersebut belum menekankan protokol anti-overfit dengan validation berdistribusi asli. "
        "Brito et al. (2024) kemudian menyoroti pentingnya explainability melalui SHAP, LIME, PDP, dan ICE agar "
        "hasil prediksi dapat diterjemahkan ke tindakan pemeliharaan yang lebih operasional. Bagi studi ini, temuan "
        "tersebut penting karena model yang dipilih tidak hanya harus akurat, tetapi juga dapat dijelaskan.",
    )
    body(
        doc,
        "Untuk dataset AI4I dan data imbalanced, beberapa studi memberi arah yang lebih spesifik. Muhidin et al. "
        "(2025) menunjukkan bahwa kombinasi Random Forest dan SMOTE efektif untuk prediksi kegagalan perangkat "
        "industri, tetapi analisisnya belum diperluas ke uji stabilitas, sensitivitas rasio oversampling, dan "
        "pengujian statistik antarmodel. Sakmar et al. (2025) melaporkan bahwa LightGBM dan CatBoost dapat menjadi "
        "pembanding yang kompetitif pada data predictive maintenance yang tidak seimbang. Dari sudut pandang "
        "operasional, Yang dan Iqbal (2025) menekankan pentingnya mempertimbangkan biaya kesalahan prediksi, "
        "sementara Taoufyq et al. (2025) merangkum bahwa pipeline PdM yang baik perlu menghubungkan deteksi, "
        "interpretasi, dan keputusan maintenance secara end-to-end.",
    )
    body(
        doc,
        "Akhirnya, Alnahhal et al. (2026) menunjukkan bahwa strategi penanganan imbalance dapat mengubah hasil "
        "klasifikasi secara signifikan, terutama pada AI4I yang memiliki distribusi kelas sangat timpang. Dale "
        "Luche et al. (2026) juga menekankan pentingnya baseline yang kuat, kalibrasi probabilitas, dan validasi "
        "yang ketat. Dengan demikian, literatur terkini mengarah pada empat kebutuhan utama: penanganan imbalance "
        "yang terkendali, pembandingan model yang adil, validasi yang tidak bias oleh data sintetis, dan "
        "interpretabilitas yang dapat ditindaklanjuti.",
    )
    caption(doc, "Tabel 1. Ringkasan penelitian terdahulu dan research gap")
    add_table(
        doc,
        ["No", "Studi", "Fokus", "Kelebihan", "Celah terhadap studi ini"],
        [
            ["1", "Çınar et al. (2020)", "Pemetaan ML untuk PdM Industri 4.0", "Landasan tren algoritma PdM", "Belum merinci pipeline anti-overfit RF pada AI4I biner"],
            ["2", "Kareem (2024)", "XGBoost vs Random Forest", "Perbandingan head-to-head ensemble", "Belum menekankan SMOTE terkendali + XAI pada AI4I"],
            ["3", "Muhidin et al. (2025)", "RF + SMOTE pada AI4I", "Fokus langsung dataset AI4I", "Evaluasi F1, McNemar, dan SHAP belum berlapis"],
            ["4", "Alnahhal et al. (2026)", "Metode imbalance PdM multiclass", "Evaluasi sistematis imbalance", "Fokus multiclass; bukan early-warning biner + SHAP"],
            ["5", "Dale Luche et al. (2026)", "Baseline robust & kalibrasi", "Validasi ketat dan probabilitas", "Belum memosisikan RF vs XGB/LR secara berlapis"],
            ["6", "Yang & Iqbal (2025)", "Perbandingan berbasis biaya", "Menyoroti batas Accuracy/F1", "Belum spesifik pipeline anti-overfit RF pada AI4I"],
            ["7", "Brito et al. (2024)", "XAI PdM (SHAP/LIME)", "Interpretabilitas operasional", "Tidak berfokus justifikasi RF pada AI4I imbalanced"],
        ],
    )
    body(
        doc,
        "Research gap. Literatur mutakhir telah membahas PdM berbasis machine learning, pembandingan Random Forest "
        "dan XGBoost, strategi penanganan imbalance, serta explainability. Namun, masih terdapat celah pada "
        "integrasi seluruh komponen tersebut dalam satu protokol yang konsisten untuk AI4I 2020, yaitu: "
        "(1) prediksi biner Machine failure dengan imbalance yang ditangani secara konservatif, bukan agresif; "
        "(2) mitigasi overfitting melalui scoring pada validation berdistribusi asli; "
        "(3) pembandingan Random Forest sebagai model usulan terhadap XGBoost, Logistic Regression–Polynomial, "
        "dan pembanding tambahan lain secara berlapis; serta "
        "(4) penguatan interpretabilitas berbasis SHAP untuk mendukung kebijakan maintenance berbasis risiko. "
        "Penelitian ini diarahkan untuk mengisi celah tersebut.",
    )


def chapter_2(doc):
    heading(doc, "Bab 2. Pemahaman Data dan Pra-Pemrosesan", 1)
    heading(doc, "2.1 Deskripsi dan Eksplorasi Data", 2)
    heading(doc, "Sumber data", 3)
    body(
        doc,
        "Dataset yang digunakan adalah AI4I 2020 Predictive Maintenance Dataset yang merepresentasikan kondisi "
        "operasional mesin milling secara realistis dan tersedia publik (Matzka, 2020). Penggunaan dataset publik "
        "dipilih karena data PdM industri riil sering terbatas akibat kerahasiaan operasional (Serradilla et al., 2022; "
        "Azari et al., 2023).",
    )
    caption(doc, "Tabel 2. Informasi dataset AI4I 2020")
    add_table(
        doc,
        ["Atribut", "Keterangan"],
        [
            ["Nama dataset", "AI4I 2020 Predictive Maintenance Classification Dataset"],
            ["Pengembang", "Stephan Matzka, HTW Berlin"],
            ["Sumber primer", "UCI Machine Learning Repository (Dataset ID: 601)"],
            ["DOI", "https://doi.org/10.24432/C5HS5C"],
            ["Lisensi", "CC BY 4.0"],
            ["Konteks industri", "Mesin milling pada lingkungan manufaktur"],
            ["Dimensi", "10.000 observasi × 14 kolom"],
            ["Missing values", "Tidak ada"],
        ],
        col_widths=[5.5, 10.5],
    )

    heading(doc, "Deskripsi fitur", 3)
    caption(doc, "Tabel 3. Deskripsi fitur dataset")
    add_table(
        doc,
        ["Fitur", "Tipe", "Peran"],
        [
            ["UDI", "Integer", "Identifier baris; dihapus"],
            ["Product ID", "String", "Identitas produk; redundan dengan Type; dihapus"],
            ["Type", "Kategorikal (L/M/H)", "Kualitas produk; di-encode ordinal 0/1/2"],
            ["Air temperature [K]", "Numerik", "Suhu udara lingkungan"],
            ["Process temperature [K]", "Numerik", "Suhu proses mesin"],
            ["Rotational speed [rpm]", "Numerik", "Kecepatan putar"],
            ["Torque [Nm]", "Numerik", "Torsi poros"],
            ["Tool wear [min]", "Numerik", "Akumulasi keausan alat"],
            ["Machine failure", "Biner", "Target primer"],
            ["TWF, HDF, PWF, OSF, RNF", "Biner", "Indikator moda kegagalan (bukan fitur prediktor)"],
        ],
    )

    heading(doc, "Kualitas data", 3)
    caption(doc, "Tabel 4. Ringkasan kualitas data")
    add_table(
        doc,
        ["Aspek", "Status", "Keterangan"],
        [
            ["Missing values", "Tidak ada", "Seluruh kolom lengkap"],
            ["Tipe data", "Konsisten", "Numerik, integer, dan kategorikal"],
            ["Imbalanced class", "Sangat ekstrem", "Normal 96,61% : Failure 3,39% (≈ 28,5 : 1)"],
            ["Kolom noninformatif", "Ada 2 kolom", "UDI dan Product ID"],
            ["Outlier", "Ada", "Terutama pada rotational speed [rpm]"],
            ["Sifat data", "Sintetis realistis", "Bukan sensor fisik nyata (Matzka, 2020)"],
        ],
    )
    body(
        doc,
        "Kekuatan dataset adalah ketiadaan missing value, ukuran yang memadai untuk eksperimen terkontrol, "
        "dan ketersediaan label moda kegagalan. Keterbatasannya adalah ketidakseimbangan ekstrem sehingga Accuracy "
        "tidak layak menjadi kriteria keputusan utama (Yang & Iqbal, 2025), serta sifat sintetis yang belum sepenuhnya "
        "mewakili noise dan drift sensor riil (Azari et al., 2023). Beberapa moda sangat jarang (RNF hanya 19 kasus) "
        "sehingga pemodelan multi-kelas memerlukan strategi khusus (Alnahhal et al., 2026).",
    )

    heading(doc, "Statistik deskriptif dan distribusi target", 3)
    caption(doc, "Tabel 5. Statistik deskriptif fitur sensor numerik")
    add_table(
        doc,
        ["Fitur", "Mean", "Min", "Max", "Std", "Pola"],
        [
            ["Air temperature [K]", "300,00", "295,3", "304,5", "2,00", "Mendekati normal"],
            ["Process temperature [K]", "310,01", "305,7", "313,8", "1,48", "Mendekati normal"],
            ["Rotational speed [rpm]", "1.538,78", "1.168", "2.886", "179,28", "Right-skewed"],
            ["Torque [Nm]", "39,99", "3,8", "76,6", "9,97", "Mendekati normal"],
            ["Tool wear [min]", "107,95", "0", "253", "63,65", "Mendekati uniform"],
        ],
    )
    caption(doc, "Tabel 6. Distribusi kelas target Machine failure")
    add_table(
        doc,
        ["Kelas", "Jumlah", "Persentase"],
        [
            ["0 – Normal", "9.661", "96,61%"],
            ["1 – Machine failure", "339", "3,39%"],
            ["Total", "10.000", "100%"],
        ],
    )
    caption(doc, "Tabel 7. Frekuensi failure mode")
    add_table(
        doc,
        ["Failure mode", "Kode", "Jumlah", "Persentase terhadap 10.000"],
        [
            ["Heat Dissipation Failure", "HDF", "115", "1,15%"],
            ["Overstrain Failure", "OSF", "98", "0,98%"],
            ["Power Failure", "PWF", "95", "0,95%"],
            ["Tool Wear Failure", "TWF", "46", "0,46%"],
            ["Random Failures", "RNF", "19", "0,19%"],
        ],
    )
    body(
        doc,
        "HDF merupakan moda paling sering, diikuti OSF dan PWF. Temuan ini berimplikasi pada kebijakan PdM: "
        "prioritas monitoring diarahkan pada pendinginan, beban mekanis, dan zona daya. RNF terlalu jarang untuk "
        "dimodelkan terpisah secara andal. Type produk didominasi L (6.000), diikuti M (2.997) dan H (1.003).",
    )

    add_figure(
        doc,
        ROOT / "fig1_target_distribution.png",
        "Gambar 1. Distribusi kelas target Machine failure dan frekuensi per failure mode",
    )
    add_figure(
        doc,
        ROOT / "fig2_sensor_distribution.png",
        "Gambar 2. Distribusi fitur sensor per kelas kegagalan (histogram + KDE)",
    )
    add_figure(
        doc,
        ROOT / "fig3_boxplot_sensor.png",
        "Gambar 3. Boxplot fitur sensor: Normal versus Machine failure",
    )
    body(
        doc,
        "Gambar 2 dan Gambar 3 menunjukkan pergeseran sebaran pada torque, tool wear, dan rotational speed. "
        "Kegagalan cenderung muncul pada torsi tinggi, keausan alat tinggi, serta kombinasi kecepatan–torsi yang "
        "menyimpang dari zona operasi normal. Pola tersebut mendukung feature engineering domain pada subbab berikutnya.",
    )
    add_figure(
        doc,
        ROOT / "fig4_correlation_heatmap.png",
        "Gambar 4. Heatmap korelasi fitur sensor terhadap Machine failure",
    )
    add_figure(
        doc,
        ROOT / "fig5_scatter_failure_zones.png",
        "Gambar 5. Scatter plot zona kegagalan: RPM versus torque (Power) dan tool wear versus torque (Strain)",
    )
    body(
        doc,
        "Korelasi linier terhadap target relatif lemah (torque ≈ 0,19; tool wear ≈ 0,11), sementara rpm dan torque "
        "berkorelasi kuat negatif (≈ −0,88). Artinya kegagalan tidak cukup dimodelkan secara linier univariat; "
        "interaksi nonlinier antar sensor lebih relevan. Gambar 5 memperlihatkan kluster kegagalan pada zona daya "
        "(rpm rendah–torsi tinggi) dan zona strain (wear tinggi–torsi tinggi), sesuai mekanisme fisik yang mendasari "
        "definisi moda PWF, OSF, dan HDF pada AI4I (Matzka, 2020).",
    )

    heading(doc, "2.2 Pra-Pemrosesan Data", 2)
    heading(doc, "Seleksi kolom dan feature engineering", 3)
    caption(doc, "Tabel 8. Kolom yang dihapus dan alasannya")
    add_table(
        doc,
        ["Kolom", "Alasan dihapus"],
        [
            ["UDI", "Hanya indeks; tidak mengandung sinyal prediktif"],
            ["Product ID", "Redundan dengan Type"],
            ["TWF, HDF, PWF, OSF, RNF", "Label moda; jika dipakai sebagai fitur terjadi target leakage"],
        ],
    )
    caption(doc, "Tabel 9. Feature engineering domain")
    add_table(
        doc,
        ["Fitur turunan", "Rumus", "Justifikasi fisik"],
        [
            ["Power", "rpm × torque", "Mendekati daya mekanis; relevan untuk PWF"],
            ["Strain", "torque × tool wear", "Beban kumulatif pada pahat; relevan untuk OSF/TWF"],
            ["Temp_diff", "process temperature − air temperature", "Gradien termal; relevan untuk HDF"],
        ],
    )
    add_figure(
        doc,
        ROOT / "fig9_kde_derived_features.png",
        "Gambar 8. KDE distribusi fitur turunan Power, Strain, dan Temp_diff per kelas",
    )
    body(
        doc,
        "Gambar 8 menunjukkan bahwa kelas failure menempati ekor distribusi Power dan Strain serta memiliki "
        "Temp_diff yang lebih terpisah. Fitur turunan ini memperkaya representasi interaksi yang tidak tertangkap "
        "oleh korelasi linier univariat.",
    )

    heading(doc, "Outlier, encoding, dan standardisasi", 3)
    body(
        doc,
        "Outlier ekstrem, terutama pada rotational speed, ditangani dengan Winsorization agar ekor distribusi "
        "tidak mendominasi pembelajaran tanpa membuang observasi kegagalan yang langka. Type di-encode secara ordinal "
        "karena terdapat urutan kualitas L < M < H. StandardScaler hanya diterapkan pada Logistic Regression karena "
        "model pohon invariant terhadap penskalaan monoton (Breiman, 2001; Chen & Guestrin, 2016).",
    )
    add_figure(
        doc,
        ROOT / "fig6_outlier_treatment.png",
        "Gambar 6. Deteksi dan penanganan outlier: sebelum versus sesudah Winsorization",
    )
    caption(doc, "Tabel 10. Encoding fitur kategorikal")
    add_table(
        doc,
        ["Fitur", "Nilai asli", "Metode", "Hasil"],
        [
            ["Type", "L, M, H", "Ordinal encoding", "L=0, M=1, H=2"],
            ["UDI / Product ID", "Identifier", "Dihapus", "—"],
            ["Fitur numerik & turunan", "Kontinu", "Tidak di-encode", "Dipakai apa adanya (tree) / di-scale (LR)"],
        ],
    )

    heading(doc, "Pembagian data dan SMOTE", 3)
    body(
        doc,
        "Split stratified 70/10/20 dengan random_state=42 menjaga proporsi kelas jarang di setiap subset "
        "(He & Garcia, 2009). SMOTE hanya diterapkan pada training untuk mencegah information leakage ke validation "
        "dan test (Chawla et al., 2002; Alnahhal et al., 2026). Rasio sampling_strategy=0,2 dipilih agar kelas failure "
        "tidak didorong ke keseimbangan artifisial 1:1, yang pada eksperimen ini justru merusak precision.",
    )
    caption(doc, "Tabel 11. Pembagian dataset train/validasi/test")
    add_table(
        doc,
        ["Split", "Proporsi", "Jumlah baris", "Peran"],
        [
            ["Training", "70%", "7.000", "Pelatihan; SMOTE hanya di sini"],
            ["Validation", "10%", "1.000", "Seleksi hiperparameter pada distribusi asli"],
            ["Test", "20%", "2.000", "Evaluasi akhir; tidak disentuh saat training/tuning"],
        ],
    )
    caption(doc, "Tabel 12. Distribusi kelas sebelum dan sesudah SMOTE pada training (strategy=0,2)")
    add_table(
        doc,
        ["Kondisi", "Normal (0)", "Failure (1)", "Rasio"],
        [
            ["Train sebelum SMOTE", "6.763", "237", "≈ 28,5 : 1"],
            ["Train sesudah SMOTE 0,2", "6.763", "≈ 1.353", "≈ 5 : 1"],
            ["Validation (tanpa SMOTE)", "≈ 966", "≈ 34", "distribusi asli"],
            ["Test (tanpa SMOTE)", "≈ 1.932", "≈ 68", "distribusi asli"],
        ],
    )
    add_figure(
        doc,
        ROOT / "fig7_smote_comparison.png",
        "Gambar 7. Perbandingan distribusi kelas training sebelum versus sesudah oversampling",
    )
    body(
        doc,
        "Model pohon dilatih pada data hasil SMOTE tanpa double-balancing (tanpa class_weight dan tanpa "
        "scale_pos_weight tambahan). Logistic Regression–Polynomial dilatih pada training asli dengan "
        "class_weight='balanced' agar baseline linear tetap peka terhadap kelas minoritas.",
    )

    heading(doc, "2.3 Penentuan Model", 2)
    add_figure(
        doc,
        ROOT / "fig10_pipeline_flowchart.png",
        "Gambar 9. Flowchart pipeline eksperimen predictive maintenance pada AI4I 2020",
        width_cm=12.5,
    )
    body(
        doc,
        "Pemilihan model disusun dalam tiga peran. Random Forest ditetapkan sebagai model usulan karena robust "
        "pada data tabular, mampu menangkap interaksi nonlinier, dan mendukung explainability berbasis pohon "
        "(Breiman, 2001; Brito et al., 2024). XGBoost menjadi pembanding utama sebagai gradient boosting yang kuat "
        "pada tabular PdM (Chen & Guestrin, 2016; Kareem, 2024). Logistic Regression–Polynomial berfungsi sebagai "
        "baseline linear agar kontribusi ensemble dapat diukur (Alnahhal et al., 2026). Gradient Boosting sklearn "
        "ditambahkan sebagai pembanding boosting default, sedangkan LightGBM/CatBoost dan model lain ditempatkan "
        "pada kelompok extended tanpa penyetelan setara.",
    )
    caption(doc, "Tabel 13. Justifikasi peran model")
    add_table(
        doc,
        ["Model", "Peran", "Alasan singkat", "Rujukan"],
        [
            ["Random Forest", "Usulan (main)", "Bagging; tabular/nonlinier; SHAP", "Breiman (2001); Kareem (2024)"],
            ["XGBoost", "Pembanding utama", "Boosting kuat + regularisasi", "Chen & Guestrin (2016)"],
            ["LR–Polynomial", "Baseline", "Acuan kinerja linear", "Alnahhal et al. (2026)"],
            ["Gradient Boosting", "Pembanding tambahan", "Boosting default sklearn", "Çınar et al. (2020)"],
            ["LightGBM & lainnya", "Extended", "Eksplorasi, belum dituning setara", "Sakmar et al. (2025)"],
        ],
    )
    caption(doc, "Tabel 14. Metrik evaluasi dan prioritas konteks manufaktur")
    add_table(
        doc,
        ["Metrik", "Prioritas", "Landasan penggunaan"],
        [
            ["F1-score @ 0,5", "Utama", "Menyeimbangkan Precision dan Recall pada data imbalanced"],
            ["Recall", "Tinggi", "Mengurangi risiko kegagalan terlewat (false negative)"],
            ["Precision", "Tinggi", "Mengendalikan false alarm perawatan"],
            ["ROC-AUC", "Pendukung", "Kualitas ranking skor risiko"],
            ["Accuracy", "Pelengkap", "Tidak dipakai sebagai kriteria keputusan utama"],
        ],
    )
    body(
        doc,
        "Dalam konteks manufaktur, biaya false negative (mesin gagal tanpa peringatan) umumnya lebih besar "
        "daripada false positive (inspeksi yang ternyata tidak perlu). Namun false alarm yang berlebihan juga "
        "menurunkan kepercayaan operator. Oleh karena itu F1@0,5 dipilih sebagai metrik keputusan, didukung Precision "
        "tinggi pada model usulan (Yang & Iqbal, 2025; Çınar et al., 2020).",
    )


def chapter_3(doc, data):
    heading(doc, "Bab 3. Pengembangan Model dan Evaluasi", 1)
    body(
        doc,
        "Bab ini menyajikan pelatihan, penyetelan hiperparameter, hasil evaluasi, serta analisis temuan. "
        "Seluruh angka merujuk pada satu rantai eksperimen hold-out yang terdokumentasi pada berkas CSV pelengkap "
        "laporan ini, bukan pencampuran antar-run.",
        first_indent=False,
    )

    heading(doc, "3.1 Proses Pelatihan Model", 2)
    body(
        doc,
        "Pelatihan mengikuti protokol anti-leakage dan anti-overfit. Validation dan test tidak di-oversample agar "
        "mencerminkan distribusi operasional. Seleksi kandidat Random Forest membandingkan pemenang GridSearch, "
        "RF dengan subsample bagging, dan soft-vote RF+GB+ExtraTrees (bobot 2-1-1) berdasarkan F1 validation. "
        "Kandidat dengan F1 validation tertinggi ditetapkan sebagai model RF final.",
    )
    caption(doc, "Tabel 15. Ringkasan protokol pelatihan")
    add_table(
        doc,
        ["Aspek", "Protokol final"],
        [
            ["Split", "Stratified 70 / 10 / 20; random_state=42"],
            ["SMOTE", "strategy=0,2; k_neighbors=5; hanya training"],
            ["Balancing pohon", "Tanpa double-balancing"],
            ["LR-Poly", "Data asli + class_weight='balanced' + StandardScaler"],
            ["Seleksi RF", "Val F1 tertinggi di antara grid / subsample / soft-vote"],
            ["Metrik keputusan", "F1@0,5 pada test; threshold val dicatat tetapi tidak menggantikan F1@0,5"],
        ],
        col_widths=[4.5, 11.5],
    )

    heading(doc, "3.2 Hyperparameter Tuning", 2)
    body(
        doc,
        "Penyetelan menggunakan GridSearchCV dengan PredefinedSplit: model dilatih pada data SMOTE dan dinilai "
        "pada validation berdistribusi asli. Pendekatan ini mengurangi seleksi model yang hanya unggul pada sampel "
        "sintetis.",
    )
    caption(doc, "Tabel 16. Hasil hyperparameter tuning")
    add_table(
        doc,
        ["Model", "Best parameters", "Catatan Val F1"],
        [
            [
                "XGBoost",
                "learning_rate=0,05; max_depth=3; min_child_weight=1; n_estimators=200; reg_lambda=1; subsample=1,0; scale_pos_weight=1",
                "0,906",
            ],
            [
                "Random Forest (grid)",
                "n_estimators=200; max_depth=None; max_features=sqrt; min_samples_split=2; min_samples_leaf=1; class_weight=None",
                "0,923 (grid) / 0,938 (final terpilih)",
            ],
            [
                "LR–Poly",
                "C=10; penalty=l1; class_weight=balanced; polynomial degree=2",
                "0,431",
            ],
        ],
    )
    ba = data["before_after"]
    ba_rows = []
    for _, r in ba.iterrows():
        ba_rows.append(
            [
                r["Model"],
                r["Split"],
                fmt3(r["F1 Before"]),
                fmt3(r["F1 After"]),
                f"{float(r['Δ F1']):+.3f}",
            ]
        )
    caption(doc, "Tabel 17. F1 sebelum versus sesudah tuning pada validation dan test")
    add_table(doc, ["Model", "Split", "F1 before", "F1 after", "Δ F1"], ba_rows)
    add_figure(
        doc,
        ROOT / "fig08_before_after_tuning.png",
        "Gambar 10. Perbandingan F1@0,5 sebelum versus sesudah hyperparameter tuning",
    )
    body(
        doc,
        "Tuning paling bermanfaat pada Random Forest (Δ F1 test +0,027 menjadi 0,880) dan sedikit pada XGBoost "
        "(+0,006 menjadi 0,846). LR-Poly naik di validation tetapi tetap underfit di test (F1 ≈ 0,403). Dengan kata lain, "
        "regularisasi dan seleksi kandidat memperbaiki generalisasi model pohon tanpa mengubah fakta bahwa baseline "
        "linear tidak mencukupi untuk deteksi seimbang.",
    )

    heading(doc, "3.3 Diagnosa Kesesuaian dan Uji Stabilitas", 2)
    caption(doc, "Tabel 18. Diagnosa gap Train versus Validation F1")
    add_table(
        doc,
        ["Model", "Train F1", "Val F1", "Gap", "Diagnosis"],
        [
            ["XGBoost (reg)", "0,9244", "0,9062", "0,0182", "Sehat (gap < 0,05)"],
            ["Random Forest (reg)", "0,9899", "0,9375", "0,0524", "Overfit ringan; Val tetap tertinggi"],
            ["LR-Poly", "0,4198", "0,4314", "−0,0116", "Underfit"],
        ],
    )
    add_figure(
        doc,
        ROOT / "fig11_learning_curves.png",
        "Gambar 11. Learning curve Train versus CV untuk XGBoost, Random Forest, dan LR-Poly",
    )
    body(
        doc,
        "XGBoost memiliki gap paling sehat. Random Forest hampir sempurna di training (kapasitas ensemble tinggi) "
        "tetapi Validation F1-nya justru tertinggi, sehingga mild overfitting tidak meniadakan generalisasi. LR-Poly "
        "gagal menangkap pola kegagalan bahkan di training, sehingga hanya layak sebagai baseline.",
    )

    cv = data["cv"]
    cv_rows = [[r["Model"], r["CV F1"], r["CV Recall"], r["CV Precision"], r["CV ROC-AUC"]] for _, r in cv.iterrows()]
    caption(doc, "Tabel 19. Stabilitas kinerja Repeated Stratified K-Fold (5×3)")
    add_table(doc, ["Model", "CV F1", "CV Recall", "CV Precision", "CV ROC-AUC"], cv_rows)
    body(
        doc,
        "Random Forest memimpin F1 (0,941 ± 0,010) dan Precision (0,982 ± 0,007) dengan simpangan baku kecil. "
        "Skor CV pada data yang telah melalui protokol SMOTE cenderung lebih optimistis dibanding test berdistribusi "
        "asli, tetapi peringkat relatif konsisten: RF > GB > XGB ≫ LR. Konsistensi ini menjadi dasar memusatkan analisis "
        "lanjutan pada Random Forest.",
    )

    heading(doc, "3.4 Evaluasi pada Test Set", 2)
    heading(doc, "3.4.1 Ranking model utama dan extended", 3)
    main = data["main"]
    main_rows = []
    for _, r in main.iterrows():
        main_rows.append(
            [
                int(r["Rank"]) if "Rank" in r else "",
                r["Model"],
                fmt3(r["Accuracy"]),
                fmt3(r["Precision"]),
                fmt3(r["Recall"]),
                fmt3(r["F1-Score"]),
                fmt3(r["ROC-AUC"]),
            ]
        )
    caption(doc, "Tabel 20. Hasil test model utama (F1@0,5)")
    add_table(
        doc,
        ["Rank", "Model", "Accuracy", "Precision", "Recall", "F1@0,5", "ROC-AUC"],
        main_rows,
    )
    ext = data["ext"]
    ext_rows = []
    for _, r in ext.iterrows():
        ext_rows.append(
            [
                int(r["Rank_Ext"]) if "Rank_Ext" in r else "",
                r["Model"],
                fmt3(r["F1-Score"]),
                fmt3(r["Recall"]),
                fmt3(r["Precision"]),
                fmt3(r["ROC-AUC"]),
            ]
        )
    caption(doc, "Tabel 21. Ranking model extended (konfigurasi bawaan, belum dituning setara)")
    add_table(doc, ["Rank", "Model", "F1", "Recall", "Precision", "ROC-AUC"], ext_rows)

    add_figure(doc, ROOT / "fig16_main_models.png", "Gambar 13. Ranking model utama pada test set")
    add_figure(doc, ROOT / "fig16_extended_models.png", "Gambar 14. Ranking model extended (default) pada test set")
    add_figure(
        doc,
        ROOT / "fig16_all_models_comparison.png",
        "Gambar 15. Perbandingan seluruh model: Random Forest unggul pada skema utama; LightGBM kompetitif di extended",
    )
    add_figure(
        doc,
        ROOT / "fig12_confusion_matrices.png",
        "Gambar 12. Confusion matrix model utama pada test set",
    )
    add_figure(
        doc,
        ROOT / "fig20_roc_main_models.png",
        "Gambar 17. Kurva ROC model utama pada test set",
    )
    body(
        doc,
        "Pada test set (2.000 baris; Failure ≈ 68), Random Forest memperoleh F1@0,5 tertinggi (0,880), Precision "
        "tertinggi (0,965), dan ROC-AUC tertinggi (0,987–0,988). Gradient Boosting sedikit unggul Recall (0,824 vs 0,809) "
        "tetapi kalah F1 dan Precision. XGBoost berada di peringkat ketiga (F1 0,846). LR-Poly mencatat Recall tinggi "
        "(0,897) dengan Precision rendah (0,260), khas baseline yang terlalu agresif memprediksi failure. LightGBM default "
        "mencapai F1 0,877—dekat dengan RF—namun belum melalui GridSearch setara sehingga tidak menggantikan ranking assignment "
        "(Kareem, 2024; Sakmar et al., 2025).",
    )

    heading(doc, "3.4.2 Sensitivitas SMOTE dan uji McNemar", 3)
    sm = data["smote"]
    sm_rows = []
    for _, r in sm.iterrows():
        sm_rows.append(
            [
                r["SMOTE_Ratio"],
                fmt3(r["Test_F1"]),
                fmt3(r["Test_Recall"]),
                fmt3(r["Test_Precision"]),
                fmt3(r["Test_ROC_AUC"]),
            ]
        )
    caption(doc, "Tabel 22. Sensitivitas rasio SMOTE pada Random Forest (komponen pohon)")
    add_table(doc, ["SMOTE ratio", "Test F1", "Recall", "Precision", "ROC-AUC"], sm_rows)
    add_figure(
        doc,
        ROOT / "fig18_smote_sensitivity_rf.png",
        "Gambar 16. Sensitivitas rasio SMOTE pada Random Forest: F1 dan Precision puncak pada rasio 0,2",
    )
    body(
        doc,
        "Menaikkan rasio SMOTE dari 0,2 ke 1,0 tidak meningkatkan F1; F1 turun dari 0,875 menjadi 0,728. Recall "
        "relatif stagnan (0,824 lalu 0,809), sementara Precision jatuh dari 0,933 ke 0,663. Oversampling agresif menambah "
        "false positive tanpa menambah kemampuan deteksi, sejalan dengan peringatan literatur bahwa metode imbalance "
        "harus dievaluasi multi-metrik (Alnahhal et al., 2026; He & Garcia, 2009).",
    )
    caption(doc, "Tabel 23. Uji McNemar (exact) pada prediksi test")
    add_table(
        doc,
        ["Perbandingan", "n01", "n10", "p-value", "Kesimpulan"],
        [
            ["RF vs Gradient Boosting", "5", "1", "0,2188", "Tidak signifikan (p ≥ 0,05)"],
            ["RF vs XGBoost", "5", "0", "0,0625", "Tidak signifikan (p ≥ 0,05)"],
        ],
    )
    body(
        doc,
        "Uji McNemar untuk sampel berpasangan menilai apakah dua pengklasifikasi berbeda pada kasus yang sama "
        "(McNemar, 1947; Dietterich, 1998). Nilai p ≥ 0,05 berarti selisih prediksi RF versus GB/XGB terlalu kecil "
        "untuk dinyatakan berbeda secara statistik. Oleh karena itu keunggulan Random Forest pada studi ini bersifat "
        "praktis (F1, Precision, ROC-AUC), bukan klaim superioritas statistik. Kehati-hatian ini justru memperkuat "
        "integritas rekomendasi karena tidak over-claiming.",
    )

    heading(doc, "3.4.3 Interpretabilitas SHAP", 3)
    shap = data["shap"]
    shap_rows = []
    for i, (feat, val) in enumerate(shap.items(), start=1):
        shap_rows.append([i, feat, f"{val:.4f}"])
    caption(doc, "Tabel 24. Ranking kepentingan fitur Random Forest berdasarkan mean |SHAP|")
    add_table(doc, ["Rank", "Fitur", "Mean |SHAP|"], shap_rows)
    add_figure(
        doc,
        ROOT / "fig22_shap_summary_rf.png",
        "Gambar 18. SHAP summary plot Random Forest pada test set",
    )
    add_figure(
        doc,
        ROOT / "fig22_shap_bar_rf.png",
        "Gambar 19. Ranking 10 fitur terpenting berdasarkan mean |SHAP|",
        width_cm=13.5,
    )
    body(
        doc,
        "SHAP menempatkan rotational speed, Strain, Temp_diff, Power, tool wear, dan torque sebagai driver utama "
        "(Lundberg & Lee, 2017). Urutan ini selaras dengan EDA dan mekanisme fisik AI4I: zona daya, strain kumulatif, "
        "dan gradien termal. Type dan process temperature berkontribusi jauh lebih kecil. Dengan demikian rumusan "
        "masalah mengenai parameter penting terjawab secara kuantitatif sekaligus dapat diterjemahkan ke prosedur "
        "monitoring perawatan (Brito et al., 2024).",
    )

    heading(doc, "3.5 Analisis Hasil dan Rekomendasi", 2)
    heading(doc, "Mengapa Random Forest direkomendasikan", 3)
    bullet(doc, "Kinerja keputusan tertinggi pada test: F1 0,880; Precision 0,965; ROC-AUC 0,987.")
    bullet(doc, "Paling stabil pada CV 5×3: F1 0,941 ± 0,010; Precision 0,982 ± 0,007.")
    bullet(doc, "Generalisasi terkendali: gap Train–Val 0,0524 (mild overfit) dengan Val F1 tertinggi.")
    bullet(doc, "Kokoh terhadap desain imbalance: puncak F1/Precision pada SMOTE 0,2, bukan pada rasio agresif.")
    bullet(doc, "Jujur secara statistik: McNemar tidak signifikan, sehingga klaim dibatasi pada keunggulan metrik operasional.")
    bullet(doc, "Dapat diinterpretasikan: SHAP selaras pengetahuan domain.")
    bullet(doc, "LightGBM default (F1 0,877) kompetitif tetapi belum dituning setara, sehingga menjadi future work, bukan pengganti ranking utama.")

    heading(doc, "Pembahasan", 3)
    body(
        doc,
        "Hasil menunjukkan bahwa prediksi Machine failure pada AI4I 2020 sangat dipengaruhi strategi imbalance "
        "dan pengendalian overfitting. Pembatasan SMOTE pada rasio 0,2 tanpa double-balancing menjaga Precision model "
        "pohon. Temuan ini selaras dengan literatur yang menekankan evaluasi multi-metrik pada PdM (Alnahhal et al., "
        "2026; Yang & Iqbal, 2025). Soft-vote RF-dominated, bila terpilih pada validation, meningkatkan Precision dengan "
        "sedikit trade-off Recall relatif terhadap Gradient Boosting. LR-Poly tetap underfit sehingga mempertegas nilai "
        "tambah ensemble nonlinier.",
    )
    body(
        doc,
        "Dari sisi operasional, HDF, OSF, dan PWF sebagai moda tersering mengarahkan kebijakan monitoring pada "
        "pendinginan, beban mekanis, dan zona daya. Fitur SHAP (rpm, Strain, Temp_diff, Power, tool wear, torque) "
        "menjadi kandidat indikator early warning. Penerapan di pabrik tetap memerlukan validasi pada data sensor riil "
        "dan kajian biaya FN–FP (Abidi et al., 2022; Dale Luche et al., 2026; Azari et al., 2023).",
    )

    heading(doc, "Keterbatasan", 3)
    bullet(doc, "Dataset bersifat sintetis sehingga pola noise/drift industri belum sepenuhnya terwakili (Matzka, 2020; Azari et al., 2023).")
    bullet(doc, "Jumlah kasus Failure pada test relatif kecil (≈ 68) sehingga metrik sensitif terhadap sedikit kesalahan prediksi.")
    bullet(doc, "Model extended belum disetel pada GridSearch yang setara dengan model utama.")
    bullet(doc, "Validation set kecil membuat optimasi threshold kurang stabil.")
    bullet(doc, "McNemar tidak signifikan; perbedaan prediksi RF vs GB/XGB tidak dapat diklaim secara statistik.")

    heading(doc, "3.6 Saran dan Pengembangan Selanjutnya", 2)
    bullet(doc, "Threshold-tuning berbasis biaya FN ≠ FP di pabrik, karena F1@0,5 adalah kriteria akademik yang dapat digeser secara operasional (Yang & Iqbal, 2025; Dale Luche et al., 2026).")
    bullet(doc, "Cost-sensitive learning atau class weight sebagai alternatif SMOTE jika rasio > 0,2 merusak precision (Alnahhal et al., 2026).")
    bullet(doc, "Penyetelan LightGBM/CatBoost dengan GridSearch setara, mengingat F1 default LightGBM sudah 0,877 (Sakmar et al., 2025; Kareem, 2024).")
    bullet(doc, "Validasi lintas mesin/waktu; CV 5×3 masih i.i.d. dan belum menjamin ketahanan terhadap drift (Azari et al., 2023; Serradilla et al., 2022).")
    bullet(doc, "Monitoring drift pada fitur SHAP utama: rpm, Strain, Temp_diff, Power, dan tool wear (Brito et al., 2024).")
    bullet(doc, "Perluasan ke klasifikasi moda kegagalan multi-label (Matzka, 2020) dan sistem end-to-end inferensi–umpan balik (Taoufyq et al., 2025; Çınar et al., 2020).")

    heading(doc, "Kesimpulan", 2)
    body(
        doc,
        "Berdasarkan rantai bukti diagnosa kesesuaian, stabilitas CV, perbandingan multi-algoritma, sensitivitas "
        "SMOTE, uji McNemar, dan interpretabilitas SHAP, Random Forest (reg) ditetapkan sebagai model rekomendasi "
        "untuk prediksi kegagalan mesin pada dataset tidak seimbang AI4I 2020. Rekomendasi operasional pada kerangka "
        "studi ini adalah Random Forest dengan threshold 0,5, SMOTE hanya pada training sebesar 0,2, pemantauan Precision, "
        "serta penjelasan fitur berbasis SHAP. Klaim keunggulan dibatasi pada metrik operasional F1/Precision/AUC, bukan "
        "pada signifikansi McNemar.",
    )


def references(doc):
    heading(doc, "Daftar Pustaka", 1)
    body(
        doc,
        "Daftar berikut disusun menurut American Psychological Association (APA, 7th ed.), diurutkan alfabetis, "
        "dengan hanging indent.",
        first_indent=False,
    )
    refs = [
        "Abidi, M. H., Mohammed, M. K., & Alkhalefah, H. (2022). Predictive maintenance planning for Industry 4.0 using machine learning for sustainable manufacturing. Sustainability, 14(6), Article 3387. https://doi.org/10.3390/su14063387",
        "Alnahhal, M., Tabash, M. I., Safi, S. K., Absy, M. S. M., & Mamadiyarov, Z. (2026). A comparative study of imbalance-handling methods in multiclass predictive maintenance. Computation, 14(4), Article 88. https://doi.org/10.3390/computation14040088",
        "American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.). https://doi.org/10.1037/0000165-000",
        "Azari, M. S., Flammini, F., Santini, S., & Caporuscio, M. (2023). A systematic literature review on transfer learning for predictive maintenance in Industry 4.0. IEEE Access, 11, 12887–12910. https://doi.org/10.1109/ACCESS.2023.3239784",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
        "Brito, L. C., Susto, G. A., Brito, J. N., & Duarte, M. A. V. (2024). Explainable predictive maintenance of rotating machines using LIME, SHAP, PDP, ICE. IEEE Access, 12, 29396–29414. https://doi.org/10.1109/ACCESS.2024.3367110",
        "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321–357. https://doi.org/10.1613/jair.953",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785–794). https://doi.org/10.1145/2939672.2939785",
        "Çınar, Z. M., Abdussalam Nuhu, A., Zeeshan, Q., Korhan, O., Asmael, M., & Safaei, B. (2020). Machine learning in predictive maintenance towards sustainable smart manufacturing in Industry 4.0. Sustainability, 12(19), Article 8211. https://doi.org/10.3390/su12198211",
        "Dale Luche, J. R., Goussain, B. G. C. dos S., & de Freitas, C. R. (2026). Robust baselines and probability calibration for TPM-oriented predictive maintenance. International Journal of Prognostics and Health Management, 17(1). https://doi.org/10.36001/ijphm.2026.v17i1.4659",
        "Dietterich, T. G. (1998). Approximate statistical tests for comparing supervised classification learning algorithms. Neural Computation, 10(7), 1895–1923. https://doi.org/10.1162/089976698300017197",
        "He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263–1284. https://doi.org/10.1109/TKDE.2008.239",
        "Kareem, A. (2024). Comparative analysis of XGBoost and random forest for predictive maintenance. Annals of the Faculty of Engineering Hunedoara, 22(4), 113–120.",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. In Advances in Neural Information Processing Systems, 30.",
        "Matzka, S. (2020). AI4I 2020 predictive maintenance dataset [Data set]. UCI Machine Learning Repository. https://doi.org/10.24432/C5HS5C",
        "McNemar, Q. (1947). Note on the sampling error of the difference between correlated proportions or percentages. Psychometrika, 12(2), 153–157. https://doi.org/10.1007/BF02295996",
        "Muhidin, A., Muhtajuddin Danny, & Surojudin, N. (2025). Prediksi kegagalan perangkat industri menggunakan random forest dan SMOTE untuk pemeliharaan preventif. Bulletin of Computer Science Research, 5(5), 1089–1094. https://doi.org/10.47065/bulletincsr.v5i5.745",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "Sakmar, M., Kadir, N. T., Shofo, P. A., & Darmawan, A. (2025). Evektivitas XGBoost LightGBM dan CatBoost pada dataset imbalanced predictive maintenance. Jurnal SINTA: Sistem Informasi dan Teknologi Komputasi, 3(1), 36–44. https://doi.org/10.61124/sinta.v3i1.145",
        "Serradilla, O., Zugasti, E., Rodriguez, J., & Zurutuza, U. (2022). Deep learning models for predictive maintenance: A survey, comparison, challenges and prospects. Applied Intelligence, 52(10), 10934–10964. https://doi.org/10.1007/s10489-021-03004-y",
        "Taoufyq, H., El Guemmat, K., Mansouri, K., & Akef, F. (2025). Predictive maintenance approaches: A systematic literature review. Journal of Industrial Engineering and Management, 18(3), 427–458. https://doi.org/10.3926/jiem.8537",
        "Yang, Y., & Iqbal, M. Z. (2025). Cost-optimised machine learning model comparison for predictive maintenance. Electronics, 14(12), Article 2497. https://doi.org/10.3390/electronics14122497",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(ref)
        set_run_font(run, size=11)

    heading(doc, "Lampiran: Sumber Reproduksi", 1)
    body(
        doc,
        "Notebook eksperimen: tugas_machinelearningS2A1.ipynb. Dataset: ai4i2020.csv (Matzka, 2020). "
        "Tabel kinerja diekspor dari final_tuned_comparison.csv, cv_repeated_stratified.csv, "
        "smote_sensitivity_rf.csv, metrics_before_after_tuning.csv, ranking_extended_models.csv, dan "
        "rf_shap_importance.csv. Gambar berasal dari keluaran notebook (fig1–fig22).",
        first_indent=False,
    )
    bullet(doc, "Kode: https://github.com/teh05/machine-learning-tr.git")
    bullet(doc, "Dataset: https://www.kaggle.com/datasets/stephanmatzka/predictive-maintenance-dataset-ai4i-2020")
    bullet(doc, "UCI: https://doi.org/10.24432/C5HS5C")


def load_data():
    main = pd.read_csv(ROOT / "final_tuned_comparison.csv")
    ext = pd.read_csv(ROOT / "final_report_ranking_extended.csv")
    cv = pd.read_csv(ROOT / "cv_repeated_stratified.csv")
    smote = pd.read_csv(ROOT / "smote_sensitivity_rf.csv")
    before_after = pd.read_csv(ROOT / "metrics_before_after_tuning.csv")
    shap = pd.read_csv(ROOT / "rf_shap_importance.csv", index_col=0).squeeze("columns")
    return {
        "main": main,
        "ext": ext,
        "cv": cv,
        "smote": smote,
        "before_after": before_after,
        "shap": shap,
    }


def build_docx():
    required = [
        "fig1_target_distribution.png",
        "fig2_sensor_distribution.png",
        "fig3_boxplot_sensor.png",
        "fig4_correlation_heatmap.png",
        "fig5_scatter_failure_zones.png",
        "fig6_outlier_treatment.png",
        "fig7_smote_comparison.png",
        "fig9_kde_derived_features.png",
        "fig10_pipeline_flowchart.png",
        "fig08_before_after_tuning.png",
        "fig11_learning_curves.png",
        "fig12_confusion_matrices.png",
        "fig16_main_models.png",
        "fig16_extended_models.png",
        "fig16_all_models_comparison.png",
        "fig18_smote_sensitivity_rf.png",
        "fig20_roc_main_models.png",
        "fig22_shap_summary_rf.png",
        "fig22_shap_bar_rf.png",
        "final_tuned_comparison.csv",
        "final_report_ranking_extended.csv",
        "cv_repeated_stratified.csv",
        "smote_sensitivity_rf.csv",
        "metrics_before_after_tuning.csv",
        "rf_shap_importance.csv",
    ]
    missing = [f for f in required if not (ROOT / f).exists()]
    if missing:
        raise FileNotFoundError("Berkas hilang: " + ", ".join(missing))

    data = load_data()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_preface_lists(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc, data)
    references(doc)
    doc.save(OUT_DOCX)
    print("DOCX saved:", OUT_DOCX)
    return OUT_DOCX


def convert_pdf(docx_path: Path, pdf_path: Path):
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(docx_path), ReadOnly=False)
        # Update TOC and PAGE fields
        try:
            word.ActiveWindow.View.Type = 3  # wdPrintView
        except Exception:
            pass
        doc.Fields.Update()
        try:
            if doc.TablesOfContents.Count:
                doc.TablesOfContents(1).Update()
        except Exception:
            pass
        doc.Repaginate()
        doc.Fields.Update()
        doc.Save()
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
        )
        doc.Close(False)
        print("PDF saved:", pdf_path)
    finally:
        word.Quit()
        pythoncom.CoUninitialize()


if __name__ == "__main__":
    docx_path = build_docx()
    convert_pdf(docx_path, OUT_PDF)
    print("DONE")
